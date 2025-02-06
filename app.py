from flask import Flask, request, jsonify
import os
import docx
import re  # Pour la détection des patterns de texte

app = Flask(__name__)

@app.route("/")
def home():
    return "API Flask est en ligne ! 🚀"

@app.route("/calcul", methods=["GET"])
def calcul():
    """ Route pour calculer le découpage optimal. """
    C = float(request.args.get("C", 10000))
    F = float(request.args.get("F", 0.7))
    W = float(request.args.get("W", 450))
    S = float(request.args.get("S", 1.6))
    I = float(request.args.get("I", 1))

    max_tokens = C * (1 - F)
    tokens_per_page = (W * S) + (I * 0.5)
    max_pages = max(1, round(max_tokens / tokens_per_page))

    return jsonify({"max_pages": max_pages})

def extract_problematic(doc_path):
    """
    Extrait la problématique d'un document DOCX.
    Si aucune problématique claire n'est trouvée, propose une reformulation basée sur le début du texte.
    """
    try:
        doc = docx.Document(doc_path)
        keywords = ["problématique", "question de recherche", "objectif de l'étude", "cette recherche s'intéresse à", "cette étude vise à"]
        found_problematic = None
        text_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            text_content.append(text)

            # Vérifie si le paragraphe contient un mot-clé lié à la problématique
            for keyword in keywords:
                if keyword.lower() in text.lower():
                    found_problematic = text
                    break  # Stoppe dès qu'on trouve une problématique

        if found_problematic:
            return f"🔍 Problématique détectée : {found_problematic}"

        # Reformulation automatique si aucune problématique claire n'est trouvée
        text_combined = " ".join(text_content[:5])  # Prend les 5 premiers paragraphes
        return f"⚠️ Problématique non trouvée. Voici une reformulation possible : {re.sub(r'[^a-zA-Z0-9 .,]', '', text_combined[:300])}..."

    except Exception as e:
        return f"❌ Erreur lors de la lecture du document : {e}"

@app.route('/detect_problematic', methods=['POST'])
def detect_problematic():
    """
    Endpoint pour détecter la problématique dans un document DOCX.
    """
    if 'file' not in request.files:
        return jsonify({"error": "Aucun fichier reçu"}), 400
    
    file = request.files['file']
    file_path = "temp.docx"
    file.save(file_path)

    problematic_text = extract_problematic(file_path)
    response = {"problematique": problematic_text}

    return jsonify(response)

def extract_sections(doc_path):
    """
    Extrait les sections importantes du document (Méthodologie, Résultats, Conclusion),
    même si les titres ne sont pas strictement "Méthodologie", "Résultats" ou "Conclusion".
    """
    try:
        doc = docx.Document(doc_path)
        sections = {
            "Méthodologie": [],
            "Résultats": [],
            "Conclusion": []
        }

        current_section = None

        for para in doc.paragraphs:
            text = para.text.strip().lower()

            # Vérifier si le paragraphe marque une nouvelle section
            if "méthodologie" in text or "méthode" in text:
                current_section = "Méthodologie"
            elif "résultats" in text or "analyse" in text or "observations" in text:
                current_section = "Résultats"
            elif "conclusion" in text or "discussion" in text:
                current_section = "Conclusion"

            # Ajouter le texte dans la bonne section
            if current_section and len(text) > 20:  # Filtrer les lignes trop courtes
                sections[current_section].append(para.text.strip())

        # Nettoyage et formatage des résultats
        for key in sections:
            if sections[key]:
                sections[key] = " ".join(sections[key][:5])  # Prend les 5 premières phrases
            else:
                sections[key] = "❌ Section non trouvée dans le document."

        return sections

    except Exception as e:
        return {"error": f"Erreur lors de l'extraction : {e}"}

@app.route('/analyze_doc', methods=['POST'])
def analyze_doc():
    """
    Analyse complète du document (Problématique, Méthodologie, Résultats, Conclusion).
    """
    if 'file' not in request.files:
        return jsonify({"error": "Aucun fichier reçu"}), 400

    file = request.files['file']
    file_path = "temp.docx"
    file.save(file_path)

    problematic_text = extract_problematic(file_path)
    sections = extract_sections(file_path)

    response = {
        "problematique": problematic_text,
        "methodologie": sections["Méthodologie"],
        "resultats": sections["Résultats"],
        "conclusion": sections["Conclusion"]
    }

    return jsonify(response)

# ✅ Vérifier que les routes sont bien chargées
print("🚀 Routes enregistrées dans Flask :")
for rule in app.url_map.iter_rules():
    print(rule)
response = {
    "problematique": problematic_text,
    "methodologie": sections.get("Méthodologie", "❌ Section Méthodologie non trouvée."),
    "resultats": sections.get("Résultats", "❌ Section Résultats non trouvée."),
    "conclusion": sections.get("Conclusion", "❌ Section Conclusion non trouvée.")
}
# ✅ Flask propre et compatible avec Render
if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))  # Port standard (5000 au lieu de 5001 pour Render)
    app.run(host="0.0.0.0", port=port)
